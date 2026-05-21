# @license
# SPDX-License-Identifier: GPL-3.0-or-later
# Copyright (c) 2024 杭州星奥传媒有限公司（影视飓风）
#
# 本文件基于原始文件（Apache-2.0 许可）进行了修改。
# 本项目基于飞书妙搭平台飓风提词器的源代码重新实现。
# 修改后按 GPL-3.0-or-later 分发。
#
import base64
from docx import Document
from docx.oxml.ns import qn


def _extract_images(doc):
    images = {}
    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            ext = rel.target_ref.split(".")[-1].lower()
        except IndexError:
            ext = "png"
        mime = {
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "gif": "image/gif",
            "bmp": "image/bmp",
            "tiff": "image/tiff",
            "svg": "image/svg+xml",
        }.get(ext, "image/png")
        b64 = base64.b64encode(rel.target_part.blob).decode("utf-8")
        images[rel.rId] = f"data:{mime};base64,{b64}"
    return images


def _inline_images_from_element(element, images):
    results = []
    for drawing in element.iter(qn("w:drawing")):
        for blip in drawing.iter(qn("a:blip")):
            r_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if r_id and r_id in images:
                results.append(
                    f'<img src="{images[r_id]}" style="max-width:100%;margin:4px 0;display:block;"/>'
                )
    return results


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _latex_escape(text: str) -> str:
    text = text.replace("\\", "\\textbackslash ")
    text = text.replace("{", "\\{")
    text = text.replace("}", "\\}")
    text = text.replace("_", "\\_")
    text = text.replace("^", "\\^{}")
    text = text.replace("$", "\\$")
    text = text.replace("#", "\\#")
    text = text.replace("%", "\\%")
    text = text.replace("&", "\\&")
    text = text.replace("~", "\\textasciitilde ")
    return text


def _brace_if_multi(text: str) -> str:
    if not text:
        return text
    if len(text) <= 1:
        return text
    return f"{{{text}}}"


def _safe_brace(text: str) -> str:
    return text if not text else f"{{{text}}}"


# ---------- OMML 公式转换（HTML 和 LaTeX，与之前相同）----------
def _omml_to_html(element) -> str:
    tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

    if tag == "t":
        return _escape(element.text or "")

    if tag in ("r", "e", "oMath", "oMathPara", "num", "den", "deg", "box",
               "dPr", "naryPr", "sub", "sup", "acc", "func", "lim", "sPre",
               "ctrlPr", "rPr", "t", "bar", "barPr"):
        return "".join(_omml_to_html(child) for child in element)

    if tag == "sSup":
        base, script = "", ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                base = _omml_to_html(child)
            elif ct == "sup":
                script = _omml_to_html(child)
        return f"{base}<sup>{script}</sup>"

    if tag == "sSub":
        base, script = "", ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                base = _omml_to_html(child)
            elif ct == "sub":
                script = _omml_to_html(child)
        return f"{base}<sub>{script}</sub>"

    if tag == "sSubSup":
        base = sub = sup = ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                base = _omml_to_html(child)
            elif ct == "sub":
                sub = _omml_to_html(child)
            elif ct == "sup":
                sup = _omml_to_html(child)
        return f"{base}<sub>{sub}</sub><sup>{sup}</sup>"

    if tag == "f":
        num_el = element.find(qn("m:num"))
        den_el = element.find(qn("m:den"))
        num = _omml_to_html(num_el) if num_el is not None else ""
        den = _omml_to_html(den_el) if den_el is not None else ""
        return f'<sup>{num}</sup>⁄<sub>{den}</sub>'

    if tag == "rad":
        deg_el = element.find(qn("m:deg"))
        e_el = element.find(qn("m:e"))
        radicand = _omml_to_html(e_el) if e_el is not None else ""
        deg = _omml_to_html(deg_el) if deg_el is not None else ""
        deg_str = f"<sup>{deg}</sup>" if deg.strip() else ""
        return f"{deg_str}√({radicand})"

    if tag == "d":
        dPr = element.find(qn("m:dPr"))
        beg_chr, end_chr = "(", ")"
        if dPr is not None:
            bc = dPr.find(qn("m:begChr"))
            if bc is not None:
                beg_chr = bc.get(qn("m:val")) or "("
            ec = dPr.find(qn("m:endChr"))
            if ec is not None:
                end_chr = ec.get(qn("m:val")) or ")"
        content = ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                content = _omml_to_html(child)
                break
        return _escape(beg_chr) + content + _escape(end_chr)

    if tag == "nary":
        naryPr = element.find(qn("m:naryPr"))
        chr_val = "∑"
        if naryPr is not None:
            chr_el = naryPr.find(qn("m:chr"))
            if chr_el is not None:
                chr_val = chr_el.get(qn("m:val")) or "∑"
        sub_el = element.find(qn("m:sub"))
        sup_el = element.find(qn("m:sup"))
        e_el = element.find(qn("m:e"))
        lower = _omml_to_html(sub_el) if sub_el is not None else ""
        upper = _omml_to_html(sup_el) if sup_el is not None else ""
        expr = _omml_to_html(e_el) if e_el is not None else ""
        parts = [_escape(chr_val)]
        if lower:
            parts.append(f"<sub>{lower}</sub>")
        if upper:
            parts.append(f"<sup>{upper}</sup>")
        parts.append(expr)
        return "".join(parts)

    if tag == "func":
        fName = element.find(qn("m:fName"))
        e_el = element.find(qn("m:e"))
        name = _omml_to_html(fName) if fName is not None else ""
        arg = _omml_to_html(e_el) if e_el is not None else ""
        return f"{name}({arg})"

    if tag == "sPre":
        sub = sup = base = ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "sub":
                sub = _omml_to_html(child)
            elif ct == "sup":
                sup = _omml_to_html(child)
            elif ct == "e":
                base = _omml_to_html(child)
        parts = []
        if sub:
            parts.append(f"<sub>{sub}</sub>")
        if sup:
            parts.append(f"<sup>{sup}</sup>")
        parts.append(base)
        return "".join(parts)

    if tag == "limLow":
        e_el = element.find(qn("m:e"))
        lim_el = element.find(qn("m:lim"))
        base = _omml_to_html(e_el) if e_el is not None else ""
        lim = _omml_to_html(lim_el) if lim_el is not None else ""
        return f"{base}<sub>{lim}</sub>"

    if tag == "limUpp":
        e_el = element.find(qn("m:e"))
        lim_el = element.find(qn("m:lim"))
        base = _omml_to_html(e_el) if e_el is not None else ""
        lim = _omml_to_html(lim_el) if lim_el is not None else ""
        return f"{base}<sup>{lim}</sup>"

    if tag == "acc":
        accPr = element.find(qn("m:accPr"))
        chr_val = "̂"
        if accPr is not None:
            chr_el = accPr.find(qn("m:chr"))
            if chr_el is not None:
                chr_val = chr_el.get(qn("m:val")) or "̂"
        e_el = element.find(qn("m:e"))
        base = _omml_to_html(e_el) if e_el is not None else ""
        return f"{base}{_escape(chr_val[:1])}"

    # 兜底：只收集直接子元素中的 m:t，避免重复
    texts = []
    for child in element:
        ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if ct == "r":
            for t in child.iter(qn("m:t")):
                if t.text:
                    texts.append(t.text)
        elif ct == "t":
            if child.text:
                texts.append(child.text)
    return _escape("".join(texts))


def _omml_to_latex(element) -> str:
    tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

    if tag == "t":
        return _latex_escape(element.text or "")

    if tag in ("r", "e", "oMath", "oMathPara", "num", "den", "deg", "box",
               "dPr", "naryPr", "sub", "sup", "acc", "func", "lim", "sPre",
               "ctrlPr", "rPr", "t", "bar", "barPr"):
        return "".join(_omml_to_latex(child) for child in element)

    if tag == "sSup":
        base = script = ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                base = _omml_to_latex(child)
            elif ct == "sup":
                script = _omml_to_latex(child)
        if not base and not script:
            return ""
        if not script:
            return base
        return f"{_brace_if_multi(base)}^{_brace_if_multi(script)}"

    if tag == "sSub":
        base = script = ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                base = _omml_to_latex(child)
            elif ct == "sub":
                script = _omml_to_latex(child)
        if not base and not script:
            return ""
        if not script:
            return base
        return f"{_brace_if_multi(base)}_{_brace_if_multi(script)}"

    if tag == "sSubSup":
        base = sub = sup = ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                base = _omml_to_latex(child)
            elif ct == "sub":
                sub = _omml_to_latex(child)
            elif ct == "sup":
                sup = _omml_to_latex(child)
        if not base and not sub and not sup:
            return ""
        parts = [_brace_if_multi(base)] if base else ["{}"]
        if sub:
            parts.append(f"_{_brace_if_multi(sub)}")
        if sup:
            parts.append(f"^{_brace_if_multi(sup)}")
        return "".join(parts)

    if tag == "f":
        num_el = element.find(qn("m:num"))
        den_el = element.find(qn("m:den"))
        num = _omml_to_latex(num_el) if num_el is not None else ""
        den = _omml_to_latex(den_el) if den_el is not None else ""
        return f"\\frac{{{num}}}{{{den}}}"

    if tag == "rad":
        deg_el = element.find(qn("m:deg"))
        e_el = element.find(qn("m:e"))
        radicand = _omml_to_latex(e_el) if e_el is not None else ""
        deg = _omml_to_latex(deg_el) if deg_el is not None else ""
        if deg.strip():
            return f"\\sqrt[{deg}]{{{radicand}}}"
        return f"\\sqrt{{{radicand}}}"

    if tag == "d":
        dPr = element.find(qn("m:dPr"))
        beg_chr, end_chr = "(", ")"
        if dPr is not None:
            bc = dPr.find(qn("m:begChr"))
            if bc is not None:
                beg_chr = bc.get(qn("m:val")) or "("
            ec = dPr.find(qn("m:endChr"))
            if ec is not None:
                end_chr = ec.get(qn("m:val")) or ")"
        content = ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                content = _omml_to_latex(child)
                break
        return f"{beg_chr}{content}{end_chr}"

    if tag == "nary":
        naryPr = element.find(qn("m:naryPr"))
        chr_val = "\\sum"
        if naryPr is not None:
            chr_el = naryPr.find(qn("m:chr"))
            if chr_el is not None:
                cv = chr_el.get(qn("m:val")) or "∑"
                cv_map = {"∑": "\\sum", "∏": "\\prod", "∫": "\\int"}
                chr_val = cv_map.get(cv, cv)
        sub_el = element.find(qn("m:sub"))
        sup_el = element.find(qn("m:sup"))
        e_el = element.find(qn("m:e"))
        lower = _omml_to_latex(sub_el) if sub_el is not None else ""
        upper = _omml_to_latex(sup_el) if sup_el is not None else ""
        expr = _omml_to_latex(e_el) if e_el is not None else ""
        parts = [chr_val]
        if lower:
            parts.append(f"_{{{lower}}}")
        if upper:
            parts.append(f"^{{{upper}}}")
        if expr:
            parts.append(f"{{{expr}}}")
        return "".join(parts)

    if tag == "func":
        fName = element.find(qn("m:fName"))
        e_el = element.find(qn("m:e"))
        name = _omml_to_latex(fName) if fName is not None else ""
        arg = _omml_to_latex(e_el) if e_el is not None else ""
        if not name:
            return arg
        return f"\\{name}{{{arg}}}"

    if tag == "sPre":
        sub = sup = base = ""
        for child in element:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "sub":
                sub = _omml_to_latex(child)
            elif ct == "sup":
                sup = _omml_to_latex(child)
            elif ct == "e":
                base = _omml_to_latex(child)
        if not sub and not sup:
            return base
        parts = ["{}"]
        if sub:
            parts.append(f"_{_brace_if_multi(sub)}")
        if sup:
            parts.append(f"^{_brace_if_multi(sup)}")
        if base:
            parts.append(_brace_if_multi(base))
        return "".join(parts)

    if tag == "limLow":
        e_el = element.find(qn("m:e"))
        lim_el = element.find(qn("m:lim"))
        base = _omml_to_latex(e_el) if e_el is not None else ""
        lim = _omml_to_latex(lim_el) if lim_el is not None else ""
        if not base:
            return f"_{_safe_brace(lim)}"
        if not lim:
            return base
        return f"{base}_{_safe_brace(lim)}"

    if tag == "limUpp":
        e_el = element.find(qn("m:e"))
        lim_el = element.find(qn("m:lim"))
        base = _omml_to_latex(e_el) if e_el is not None else ""
        lim = _omml_to_latex(lim_el) if lim_el is not None else ""
        if not base:
            return f"^{_safe_brace(lim)}"
        if not lim:
            return base
        return f"{base}^{_safe_brace(lim)}"

    if tag == "acc":
        accPr = element.find(qn("m:accPr"))
        chr_val = "̂"
        if accPr is not None:
            chr_el = accPr.find(qn("m:chr"))
            if chr_el is not None:
                chr_val = chr_el.get(qn("m:val")) or "̂"
        e_el = element.find(qn("m:e"))
        base = _omml_to_latex(e_el) if e_el is not None else ""
        accent_map = {
            "̂": "hat", "̃": "tilde", "̄": "bar", "̇": "dot",
            "̈": "ddot", "⃗": "vec", "̂": "widehat", "̃": "widetilde",
            "́": "acute", "̀": "grave", "̌": "check", "̆": "breve",
            "→": "overrightarrow", "←": "overleftarrow",
            "⇀": "overrightharpoon", "↼": "overleftharpoon",
        }
        cmd = accent_map.get(chr_val, "hat")
        return f"\\{cmd}{{{base}}}"

    if tag == "bar":
        e_el = element.find(qn("m:e"))
        base = _omml_to_latex(e_el) if e_el is not None else ""
        pos_el = element.find(qn("m:barPr"))
        pos = "top"
        if pos_el is not None:
            pos_tag = pos_el.find(qn("m:pos"))
            if pos_tag is not None:
                pos = (pos_tag.get(qn("m:val")) or "top").lower()
        if pos == "bot":
            return f"\\underline{{{base}}}"
        return f"\\overline{{{base}}}"

    texts = []
    for child in element:
        ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if ct == "r":
            for t in child.iter(qn("m:t")):
                if t.text:
                    texts.append(_latex_escape(t.text))
        elif ct == "t":
            if child.text:
                texts.append(_latex_escape(child.text))
    return "".join(texts)


def _has_formulas(para) -> bool:
    for math_tag in (qn("m:oMath"), qn("m:oMathPara")):
        for _ in para._element.iter(math_tag):
            return True
    return False


def _convert_wr_element(wr_element, images, formula_mode="html") -> list[str]:
    rPr = wr_element.find(qn("w:rPr"))
    bold = italic = underline = superscript = subscript = False
    if rPr is not None:
        bold = rPr.find(qn("w:b")) is not None
        italic = rPr.find(qn("w:i")) is not None
        underline = rPr.find(qn("w:u")) is not None
        vertAlign = rPr.find(qn("w:vertAlign"))
        if vertAlign is not None:
            val = vertAlign.get(qn("w:val"))
            if val == "superscript":
                superscript = True
            elif val == "subscript":
                subscript = True

    parts = []
    current_texts = []

    def flush_text():
        nonlocal current_texts
        if current_texts:
            txt = "".join(current_texts)
            current_texts = []
            if txt:
                txt = _escape(txt)
                if subscript:
                    txt = f"<sub>{txt}</sub>"
                elif superscript:
                    txt = f"<sup>{txt}</sup>"
                if bold:
                    txt = f"<strong>{txt}</strong>"
                if italic:
                    txt = f"<em>{txt}</em>"
                if underline:
                    txt = f"<u>{txt}</u>"
                parts.append(txt)

    for child in wr_element:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "t":
            flush_text()
            current_texts.append(child.text or "")
        elif tag == "br":
            flush_text()
            parts.append("<br>")
        elif tag == "drawing":
            flush_text()
            parts.extend(_inline_images_from_element(wr_element, images))
        elif tag in ("oMath", "oMathPara"):
            flush_text()
            if formula_mode == "latex":
                latex = _omml_to_latex(child)
                if latex:
                    escaped = latex.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    if tag == "oMathPara":
                        parts.append(f"$${escaped}$$")
                    else:
                        parts.append(f"${escaped}$")
            else:
                html = _omml_to_html(child)
                if html:
                    parts.append(f'<span class="formula">{html}</span>')

    flush_text()
    return parts


def _para_to_html(para, images, formula_mode="html"):
    parts = []
    # 直接遍历段落的子元素，按文档原始顺序处理
    for child in para._element:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            # 文本 run（可能内含公式、图片、上下标）
            parts.extend(_convert_wr_element(child, images, formula_mode))
        elif tag in ("oMath", "oMathPara"):
            # 段落级公式（独立成段的显示公式或内联公式）
            if formula_mode == "latex":
                latex = _omml_to_latex(child)
                if latex:
                    escaped = latex.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    if tag == "oMathPara":
                        parts.append(f"$${escaped}$$")
                    else:
                        parts.append(f"${escaped}$")
            else:
                html = _omml_to_html(child)
                if html:
                    parts.append(f'<span class="formula">{html}</span>')
        # 忽略其他元素（如属性标记等）

    joined = "".join(parts)
    if not joined.strip():
        return ""
    style_name = para.style.name.lower() if para.style else ""
    if style_name.startswith("heading 1"):
        return f"<h1>{joined}</h1>"
    elif style_name.startswith("heading 2"):
        return f"<h2>{joined}</h2>"
    elif style_name.startswith("heading 3"):
        return f"<h3>{joined}</h3>"
    return f"<p>{joined}</p>"


def _cell_to_html(cell, formula_mode="html") -> str:
    parts = []
    for para in cell.paragraphs:
        html = _para_to_html(para, {}, formula_mode)
        if html:
            parts.append(html)
    return "".join(parts) if parts else _escape(cell.text.strip())


def _table_to_html(table, formula_mode="html"):
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_html = _cell_to_html(cell, formula_mode)
            cells.append(f"<td>{cell_html}</td>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return (
        '<table border="1" style="border-collapse:collapse;width:100%;margin:8px 0;">'
        + "".join(rows)
        + "</table>"
    )


def _walk_body(element, paras, tables, images, formula_mode):
    html_parts = []
    has_formulas = False
    for child in element:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            para = paras.get(id(child))
            if para:
                if _has_formulas(para):
                    has_formulas = True
                h = _para_to_html(para, images, formula_mode)
                if h:
                    html_parts.append(h)
        elif tag == "tbl":
            table = tables.get(id(child))
            if table:
                html_parts.append(_table_to_html(table, formula_mode))
        elif tag in ("sdt", "ins", "del", "smartTag", "customXml",
                      "moveFrom", "moveTo", "subDoc"):
            sub_parts, sub_formulas = _walk_body(child, paras, tables, images, formula_mode)
            html_parts.extend(sub_parts)
            has_formulas = has_formulas or sub_formulas
    return html_parts, has_formulas


def import_docx_file(filepath: str, formula_mode: str = "html") -> tuple[str, bool]:
    doc = Document(filepath)
    images = _extract_images(doc)

    paras = {id(p._element): p for p in doc.paragraphs}
    tables = {id(t._element): t for t in doc.tables}
    body = doc.element.body

    html_parts, has_formulas = _walk_body(body, paras, tables, images, formula_mode)
    return "\n".join(html_parts), has_formulas